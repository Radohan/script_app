import os
import re
import traceback
import pandas as pd
import tempfile
from PyQt5.QtWidgets import QFileDialog, QMessageBox
import functools


class DocumentParser:
    """
    Handles parsing Word and PDF documents to extract tables and match content
    with MXLIFF translations.
    """

    def __init__(self, parent=None):
        """Initialize with parent widget for dialogs."""
        self.parent = parent
        self.tables = []
        self.file_path = None
        self.file_type = None
        self.debug_info = []  # Store debug info for troubleshooting

    def parse_document(self, file_path=None):
        """Parse document with improved performance."""
        if file_path:
            self.file_path = file_path

        if not self.file_path:
            return False

        try:
            # Determine file type and use appropriate parser
            if self.file_path.lower().endswith('.docx'):
                return self._parse_docx_optimized()
            elif self.file_path.lower().endswith('.pdf'):
                return self._parse_pdf_optimized()
            else:
                return False
        except Exception as e:
            self.parent.log(f"Error parsing document: {str(e)}")
            import traceback
            self.parent.log(traceback.format_exc())
            return False

    def _parse_docx_optimized(self):
        """Parse DOCX file with optimized performance."""
        import docx

        try:
            self.doc = docx.Document(self.file_path)
            self.tables = []

            # Process tables in chunks
            table_count = len(self.doc.tables)
            for i, table in enumerate(self.doc.tables):
                # Report progress to parent if available
                if hasattr(self.parent, 'worker') and hasattr(self.parent.worker, 'progress_signal'):
                    progress = int(20 + (i / table_count) * 30)  # 20-50% progress range
                    self.parent.worker.progress_signal.emit(
                        progress,
                        f"Processing table {i + 1} of {table_count}..."
                    )

                # Allow event processing periodically
                if i % 5 == 0:
                    QApplication.processEvents()

                # Process table data as before
                # ...

            return True

        except Exception as e:
            self.parent.log(f"Error parsing DOCX: {str(e)}")
            import traceback
            self.parent.log(traceback.format_exc())
            return False

    def select_document(self):
        """
        Show a file dialog to select a Word or PDF document.

        Returns:
            str: Path to the selected file or None if canceled
        """
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getOpenFileName(
            self.parent,
            "Select Document File",
            "",
            "Document Files (*.docx *.pdf);;Word Documents (*.docx);;PDF Files (*.pdf);;All Files (*)",
            options=options
        )

        if file_path:
            self.file_path = file_path
            _, extension = os.path.splitext(file_path)
            self.file_type = extension.lower()
            return file_path

        return None

    def log_debug(self, message):
        """Add debug info for troubleshooting."""
        print(f"DocumentParser: {message}")
        self.debug_info.append(message)

    def parse_document(self):
        """
        Parse the selected document and extract tables.

        Returns:
            bool: True if parsing succeeded, False otherwise
        """
        if not self.file_path:
            return False

        self.debug_info = []  # Reset debug info
        self.tables = []  # Reset tables

        try:
            if self.file_type == '.docx':
                try:
                    # Try importing docx - if it fails, provide helpful error
                    import docx
                    success = self._parse_word_document()
                except ImportError:
                    self.log_debug("python-docx package not installed")
                    if self.parent:
                        QMessageBox.critical(
                            self.parent,
                            "Missing Dependency",
                            "The python-docx package is required to parse Word documents.\n"
                            "Please install it with: pip install python-docx"
                        )
                    return False
            elif self.file_type == '.pdf':
                try:
                    # Try importing PyMuPDF - if it fails, provide helpful error
                    import fitz
                    success = self._parse_pdf_document()
                except ImportError:
                    self.log_debug("PyMuPDF package not installed")
                    if self.parent:
                        QMessageBox.critical(
                            self.parent,
                            "Missing Dependency",
                            "The PyMuPDF package is required to parse PDF documents.\n"
                            "Please install it with: pip install PyMuPDF"
                        )
                    return False
            else:
                if self.parent:
                    QMessageBox.warning(
                        self.parent,
                        "Unsupported File Type",
                        f"The file type {self.file_type} is not supported. Please select a .docx or .pdf file."
                    )
                return False

            # Check if we found any tables
            if not self.tables:
                if self.parent:
                    QMessageBox.warning(
                        self.parent,
                        "No Tables Found",
                        "No tables were found in the document."
                    )
                return False

            # Try direct table detection - if none have Conversation column, perform fallback method
            conversation_tables = self.get_conversation_tables()

            if not conversation_tables:
                self.log_debug("No tables with 'Conversation' column found, trying fallback detection")
                # Fallback: Check for Table column that might contain Conversation values
                self.detect_conversation_tables_fallback()
                conversation_tables = self.get_conversation_tables()

            if not conversation_tables:
                # Show debug info in the message box to help troubleshoot
                table_info = []
                for i, table in enumerate(self.tables):
                    cols = list(table['dataframe'].columns)
                    table_info.append(f"Table {i + 1}: Columns = {cols}")

                debug_str = "\n".join(table_info)

                if self.parent:
                    QMessageBox.warning(
                        self.parent,
                        "No Conversation Tables Found",
                        f"The document does not contain any tables with a 'Conversation' column or with conversation data.\n\n"
                        f"Debug info:\n{debug_str}"
                    )
                return False

            self.log_debug(f"Final count of conversation tables: {len(conversation_tables)}")
            return True

        except Exception as e:
            self.log_debug(f"Error parsing document: {str(e)}")
            self.log_debug(traceback.format_exc())
            if self.parent:
                QMessageBox.critical(
                    self.parent,
                    "Document Parsing Error",
                    f"An error occurred while parsing the document: {str(e)}\n\n"
                    f"Debug info:\n{chr(10).join(self.debug_info)}"
                )
            return False

    def _parse_word_document(self):
        """Parse tables from a Word document."""
        import docx
        doc = docx.Document(self.file_path)

        # Log document info
        self.log_debug(f"Word document contains {len(doc.tables)} tables")

        # Extract tables from Word document
        for i, table in enumerate(doc.tables):
            # Get number of rows and columns
            rows = len(table.rows)
            cols = len(table.rows[0].cells) if rows > 0 else 0

            self.log_debug(f"Table {i + 1}: {rows} rows, {cols} columns")

            if rows <= 1:  # Skip tables with only header or no data
                self.log_debug(f"Skipping Table {i + 1} (not enough rows)")
                continue

            # Convert table to a list of lists
            data = []
            headers = []

            # Get headers from first row
            for cell in table.rows[0].cells:
                header_text = self._clean_text(cell.text)
                self.log_debug(f"Header found: '{header_text}'")
                headers.append(header_text)

            # Get data from remaining rows
            for row in table.rows[1:]:
                row_data = []
                for cell in row.cells:
                    row_data.append(self._clean_text(cell.text))
                data.append(row_data)

            # Convert to DataFrame
            if headers and data:
                # Make sure data rows have same length as headers
                uniform_data = []
                for row in data:
                    # Pad or truncate row to match headers length
                    if len(row) < len(headers):
                        row = row + [''] * (len(headers) - len(row))
                    elif len(row) > len(headers):
                        row = row[:len(headers)]
                    uniform_data.append(row)

                df = pd.DataFrame(uniform_data, columns=headers)

                # Check for "Conversation" or similar column (case insensitive)
                has_conversation = False
                conversation_col = None

                for col in df.columns:
                    # Check if column name contains "conversation" (case insensitive)
                    if 'conversation' in col.lower() or col.lower() == 'conv' or col.lower() == 'table':
                        has_conversation = True
                        conversation_col = col
                        self.log_debug(f"Found conversation column: '{col}'")
                        break

                self.tables.append({
                    'id': i + 1,
                    'dataframe': df,
                    'has_conversation_column': has_conversation,
                    'conversation_column': conversation_col
                })

                self.log_debug(f"Added Table {i + 1} with {len(df)} rows, {len(df.columns)} columns")

        return len(self.tables) > 0

    def _parse_pdf_document(self):
        """Parse tables from a PDF document."""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(self.file_path)
            table_id = 1

            self.log_debug(f"PDF document contains {len(doc)} pages")

            # First try to extract tables using PyMuPDF's table finder
            try:
                self._extract_tables_with_pymupdf(doc, table_id)
            except Exception as e:
                self.log_debug(f"PyMuPDF table extraction failed: {str(e)}")
                self.log_debug("Falling back to text extraction to find tables")
                self._extract_tables_from_text(doc, table_id)

            return len(self.tables) > 0
        except Exception as e:
            self.log_debug(f"Error in PDF parsing: {str(e)}")
            self.log_debug(traceback.format_exc())
            raise

    def _extract_tables_with_pymupdf(self, doc, table_id):
        """Extract tables using PyMuPDF's table finder."""
        import fitz
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)

            # Extract tables using PyMuPDF's table extractor
            tables = page.find_tables()
            self.log_debug(f"Page {page_num + 1}: Found {len(tables.tables) if tables else 0} tables")

            if not tables or not tables.tables:
                continue

            for table in tables.tables:
                self.log_debug(f"Processing table with {table.row_count} rows, {table.cols} columns")

                if table.row_count <= 1:  # Skip tables with only header or no data
                    self.log_debug("Skipping table (not enough rows)")
                    continue

                headers = []
                data = []

                # Get headers from first row
                if table.row_count > 0:
                    for col_idx in range(table.cols):
                        cell_idx = col_idx
                        if cell_idx < len(table.cells) and table.cells[cell_idx]:
                            rect = table.cells[cell_idx].rect
                            text = page.get_text("text", clip=rect).strip()
                            headers.append(self._clean_text(text))
                            self.log_debug(f"Header found: '{text}'")
                        else:
                            headers.append(f"Column_{col_idx + 1}")

                # Get data from remaining rows
                for row_idx in range(1, table.row_count):
                    row_data = []
                    for col_idx in range(table.cols):
                        cell_idx = row_idx * table.cols + col_idx
                        if cell_idx < len(table.cells) and table.cells[cell_idx]:
                            rect = table.cells[cell_idx].rect
                            text = page.get_text("text", clip=rect).strip()
                            row_data.append(self._clean_text(text))
                        else:
                            row_data.append("")
                    data.append(row_data)

                # Create DataFrame
                if headers and data:
                    # Process the table data into a DataFrame
                    df = self._create_dataframe_from_table_data(headers, data, table_id)
                    if df is not None:
                        table_id += 1

    def _extract_tables_from_text(self, doc, table_id):
        """Extract tables by analyzing text content for table structures."""
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()

            # Look for table patterns (e.g., text with consistent spacing that looks like a table)
            # This is a simplified approach - could be enhanced with more sophisticated table detection
            lines = text.split('\n')

            # Look for potential table header rows - lines with multiple "|" or tab characters
            table_start_indices = []
            for i, line in enumerate(lines):
                if line.count('|') > 2 or line.count('\t') > 2:
                    table_start_indices.append(i)

            # Process each potential table
            for start_idx in table_start_indices:
                # Look for the end of the table
                end_idx = start_idx + 1
                while end_idx < len(lines) and (lines[end_idx].count('|') > 1 or lines[end_idx].count('\t') > 1):
                    end_idx += 1

                # If the table has at least a header and one data row
                if end_idx - start_idx > 1:
                    table_lines = lines[start_idx:end_idx]

                    # Determine the delimiter (| or \t)
                    delimiter = '|' if '|' in table_lines[0] else '\t'

                    # Parse the header and data
                    headers = [col.strip() for col in table_lines[0].split(delimiter) if col.strip()]
                    data = []
                    for line in table_lines[1:]:
                        row_data = [col.strip() for col in line.split(delimiter)]
                        # Ensure row has same number of columns as header
                        if len(row_data) < len(headers):
                            row_data.extend([''] * (len(headers) - len(row_data)))
                        elif len(row_data) > len(headers):
                            row_data = row_data[:len(headers)]
                        data.append(row_data)

                    # Create DataFrame
                    if headers and data:
                        df = self._create_dataframe_from_table_data(headers, data, table_id)
                        if df is not None:
                            table_id += 1

    def _create_dataframe_from_table_data(self, headers, data, table_id):
        """Create a DataFrame from table headers and data."""
        # Handle case where headers might be empty or short
        while len(headers) < max(len(row) for row in data if row):
            headers.append(f"Column_{len(headers) + 1}")

        # Make sure data rows have same length as headers
        uniform_data = []
        for row in data:
            # Pad or truncate row to match headers length
            if len(row) < len(headers):
                row = row + [''] * (len(headers) - len(row))
            elif len(row) > len(headers):
                row = row[:len(headers)]
            uniform_data.append(row)

        # Create DataFrame
        try:
            df = pd.DataFrame(uniform_data, columns=headers)

            # Check for "Conversation" or similar column (case insensitive)
            has_conversation = False
            conversation_col = None

            for col in df.columns:
                # Check if column name contains "conversation" (case insensitive)
                col_lower = str(col).lower()
                if 'conversation' in col_lower or col_lower == 'conv' or col_lower == 'table':
                    has_conversation = True
                    conversation_col = col
                    self.log_debug(f"Found conversation column: '{col}'")
                    break

            self.tables.append({
                'id': table_id,
                'dataframe': df,
                'has_conversation_column': has_conversation,
                'conversation_column': conversation_col
            })

            self.log_debug(f"Added Table {table_id} with {len(df)} rows, {len(df.columns)} columns")
            return df
        except Exception as e:
            self.log_debug(f"Error creating DataFrame: {str(e)}")
            return None

    def detect_conversation_tables_fallback(self):
        """
        Check for tables that have a 'Table' column with 'Conversation' values
        or other indicators that the table contains conversation data.
        """
        for table in self.tables:
            if table['has_conversation_column']:
                continue  # Already identified as a conversation table

            df = table['dataframe']

            # Check for any column that might be the 'table type' column
            for col in df.columns:
                if 'table' in str(col).lower():
                    # Check if this column contains any 'conversation' values
                    values = df[col].astype(str).str.lower()
                    has_conversation_values = values.str.contains('conversation').any()

                    if has_conversation_values:
                        self.log_debug(f"Found table with '{col}' column containing 'conversation' values")
                        table['has_conversation_column'] = True
                        table['conversation_column'] = col
                        break

            # If still not found, look for a column with 'Conversation' in most of its values
            if not table['has_conversation_column']:
                for col in df.columns:
                    # Check each string value in this column
                    values = df[col].astype(str).str.lower()
                    conversation_count = values.str.contains('conversation').sum()

                    # If 'conversation' appears in many of the values, this might be our column
                    if conversation_count > len(df) * 0.5:  # More than 50% have 'conversation'
                        self.log_debug(f"Found table with column '{col}' containing mostly 'conversation' values")
                        table['has_conversation_column'] = True
                        table['conversation_column'] = col
                        break

    def get_conversation_tables(self):
        """
        Get tables that have a 'Conversation' column.

        Returns:
            list: List of dataframes that have a Conversation column
        """
        return [t for t in self.tables if t['has_conversation_column']]

    def match_content_with_mxliff(self, mxliff_data):
        """
        Match content from tables with source text in MXLIFF data.

        Args:
            mxliff_data (list): List of processed data from MXLIFF parser

        Returns:
            dict: Dictionary with matches and updates to apply
        """
        import time
        import os
        from difflib import SequenceMatcher

        start_time = time.time()
        debug_log_path = os.path.join(os.path.dirname(__file__), 'match_debug.log')

        with open(debug_log_path, 'w', encoding='utf-8') as debug_log:
            debug_log.write(f"Matching started at: {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
            debug_log.write("=" * 50 + "\n")

            # Ensure we have tables to process
            if not self.tables:
                debug_log.write("No tables found in the document.\n")
                return {'matches': 0, 'updates': []}

            # Get conversation tables
            conversation_tables = self.get_conversation_tables()
            if not conversation_tables:
                debug_log.write("No conversation tables found.\n")
                return {'matches': 0, 'updates': []}

            # Preprocess MXLIFF items
            mxliff_items = {}
            for data in mxliff_data:
                if not data.get('is_header', True) and 'item' in data:
                    item = data['item']
                    source_text = item.get('source_text', '').strip()
                    if source_text:
                        # Use cleaned text as key for more robust matching
                        clean_source = self._clean_text_for_comparison(source_text)
                        mxliff_items[clean_source] = {
                            'original_text': source_text,
                            'item': item
                        }

            debug_log.write(f"Total MXLIFF items: {len(mxliff_items)}\n")
            debug_log.write(f"Conversation tables: {len(conversation_tables)}\n")

            updates = []
            match_count = 0

            # Iterate through conversation tables
            for table_index, table in enumerate(conversation_tables, 1):
                df = table['dataframe']
                debug_log.write(f"\nProcessing Table {table_index}:\n")

                # Find Source and COTeam Comments columns
                source_col = None
                comments_col = None

                for col in df.columns:
                    col_str = str(col).lower()
                    if 'source' in col_str:
                        source_col = col
                    elif ('coteam' in col_str and 'comment' in col_str) or (
                            'comment' in col_str and 'coteam' in col_str):
                        comments_col = col

                # Fallback comments column detection
                if not comments_col:
                    comments_col = next((col for col in df.columns if 'comment' in str(col).lower()), None)

                if not source_col or not comments_col:
                    debug_log.write("Skipping table - required columns not found\n")
                    continue

                # Batch process the table rows
                for row_index, row in df.iterrows():
                    source_text = str(row[source_col]).strip()
                    comment_text = str(row[comments_col]).strip()

                    # Skip empty entries
                    if not source_text or not comment_text:
                        continue

                    # Clean source text
                    clean_source = self._clean_text_for_comparison(source_text)

                    # Exact match
                    if clean_source in mxliff_items:
                        match_data = mxliff_items[clean_source]
                        update_entry = {
                            'key': match_data['item'].get('key', ''),
                            'source_text': clean_source,
                            'comment': comment_text,
                            'match_type': 'exact'
                        }
                        updates.append(update_entry)
                        match_count += 1
                        debug_log.write(f"Exact Match: {clean_source[:100]}...\n")
                        continue

                    # Fuzzy matching with optimization
                    best_match = None
                    best_ratio = 0.8  # Similarity threshold

                    for mxliff_source, mxliff_data in mxliff_items.items():
                        # Skip very short strings
                        if len(clean_source) < 10 or len(mxliff_source) < 10:
                            continue

                        # Use SequenceMatcher for more efficient similarity calculation
                        ratio = SequenceMatcher(None, clean_source, mxliff_source).ratio()

                        if ratio > best_ratio:
                            best_ratio = ratio
                            best_match = mxliff_data

                    if best_match:
                        update_entry = {
                            'key': best_match['item'].get('key', ''),
                            'source_text': best_match['original_text'],
                            'comment': comment_text,
                            'match_type': 'fuzzy',
                            'match_ratio': best_ratio
                        }
                        updates.append(update_entry)
                        match_count += 1
                        debug_log.write(f"Fuzzy Match: {clean_source[:100]}... (Ratio: {best_ratio})\n")

            # Final logging
            debug_log.write("\n" + "=" * 50 + "\n")
            debug_log.write(f"Total matches found: {match_count}\n")
            debug_log.write(f"Total time taken: {time.time() - start_time:.2f} seconds\n")

        print(f"Matching completed. Debug log saved to {debug_log_path}")

        return {
            'matches': match_count,
            'updates': updates
        }



    @functools.lru_cache(maxsize=1000)
    def _clean_text_for_comparison(self, text):
        """Cached and optimized text cleaning with memoization."""
        if not text:
            return ""

        text = str(text).strip()

        try:
            text = re.sub(r'[''‚‛""„‟\x00-\x1F\x7F]', lambda m: {
                ''': "'", ''': "'",
                '"': '"', '"': '"'
            }.get(m.group(0), ' '), text)
        except Exception as e:
            print(f"Error cleaning text: {e}")
            return text

        return re.sub(r'\s+', ' ', text).strip()

    def _clean_text_for_comparison(self, text):
        """
        Clean and normalize text for comparison purposes.

        Args:
            text (str): Input text to clean

        Returns:
            str: Cleaned and normalized text
        """
        # Handle None or empty input
        if text is None:
            return ""

        # Convert to string and strip whitespace
        text = str(text).strip()

        # Handle empty string after stripping
        if not text:
            return ""

        try:
            # Replace various quote characters and normalize
            text = re.sub(r'[''‚‛""„‟\x00-\x1F\x7F]', lambda m: {
                ''': "'",   # Left single quote
                ''': "'",  # Right single quote
                '"': '"',  # Left double quote
                '"': '"',  # Right double quote
            }.get(m.group(0), ' '), text)

            # Normalize whitespace
            text = re.sub(r'\s+', ' ', text).strip()
        except Exception as e:
            print(f"Error cleaning text: {e}")
            return text

        return text

    def _clean_text(self, text):
        """
        Basic text cleaning method.

        Args:
            text (str): Input text to clean

        Returns:
            str: Cleaned text
        """
        if text is None:
            return ""

        # Convert to string
        text = str(text)

        # Replace non-breaking spaces and other problematic whitespace
        text = text.replace('\xa0', ' ')

        # Remove control characters
        text = ''.join(c if ord(c) >= 32 or c in '\n\r\t' else ' ' for c in text)

        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text).strip()

        return text

    # Ensure other methods that might use these cleaning methods are present
    def match_content_with_mxliff(self, mxliff_data):
        import time
        import os
        from difflib import SequenceMatcher

        start_time = time.time()

        # Early quick checks
        if not mxliff_data or not self.tables:
            return {'matches': 0, 'updates': []}

        # Efficient preprocessing
        def preprocess_text(text):
            return self._clean_text_for_comparison(str(text).strip())

        # Create efficient lookup structures
        mxliff_lookup = {}
        for data in mxliff_data:
            if not data.get('is_header', True) and 'item' in data:
                item = data['item']
                source_text = str(item.get('source_text', '')).strip()
                if source_text:
                    clean_source = preprocess_text(source_text)
                    if len(clean_source) >= 10:  # Filter very short texts
                        mxliff_lookup[clean_source] = {
                            'key': item.get('key', ''),
                            'original_text': source_text
                        }

        # Get conversation tables with early filtering
        conversation_tables = [
            table for table in self.get_conversation_tables()
            if len(table['dataframe']) > 0
        ]

        if not conversation_tables:
            return {'matches': 0, 'updates': []}

        # Optimized matching with early stopping
        updates = []
        match_count = 0
        MAX_MATCHES = 500  # Prevent unlimited matches

        for table in conversation_tables:
            df = table['dataframe']

            # Smart column detection
            source_col = next((col for col in df.columns if 'source' in str(col).lower()), None)
            comments_col = next((col for col in df.columns if 'comment' in str(col).lower()), None)

            if not source_col or not comments_col:
                continue

            for _, row in df.iterrows():
                if match_count >= MAX_MATCHES:
                    break

                source_text = str(row[source_col]).strip()
                comment_text = str(row[comments_col]).strip()

                if not source_text or not comment_text:
                    continue

                clean_source = preprocess_text(source_text)

                # Exact match
                if clean_source in mxliff_lookup:
                    match_data = mxliff_lookup[clean_source]
                    updates.append({
                        'key': match_data['key'],
                        'source_text': clean_source,
                        'comment': comment_text,
                        'match_type': 'exact'
                    })
                    match_count += 1
                    continue

                # Fuzzy matching with early stopping
                best_match = None
                best_ratio = 0.8  # Similarity threshold

                for mxliff_source, mxliff_data in mxliff_lookup.items():
                    # Skip very short strings and improve matching efficiency
                    if len(clean_source) < 10 or len(mxliff_source) < 10:
                        continue

                    # Quick length-based filtering before detailed comparison
                    if abs(len(clean_source) - len(mxliff_source)) > 5:
                        continue

                    # Faster similarity calculation
                    ratio = SequenceMatcher(None, clean_source, mxliff_source).ratio()

                    if ratio > best_ratio:
                        best_ratio = ratio
                        best_match = mxliff_data

                if best_match:
                    updates.append({
                        'key': best_match['key'],
                        'source_text': best_match['original_text'],
                        'comment': comment_text,
                        'match_type': 'fuzzy',
                        'match_ratio': best_ratio
                    })
                    match_count += 1

        # Performance logging
        print(f"Matching completed: {match_count} matches in {time.time() - start_time:.2f} seconds")

        return {
            'matches': match_count,
            'updates': updates
        }

    def _similarity_ratio(self, str1, str2):
        """Calculate similarity ratio between two strings."""
        # Simple implementation using longest common subsequence
        import difflib
        return difflib.SequenceMatcher(None, str1, str2).ratio()
